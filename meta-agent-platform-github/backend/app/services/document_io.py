from __future__ import annotations

import html
import re
import uuid
import zipfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from fastapi import HTTPException, UploadFile

from ..config import (
    ARTIFACT_STORAGE_DIR,
    MAX_WORKFLOW_UPLOAD_BYTES,
    UPLOAD_STORAGE_DIR,
)
from ..database import get_connection

SUPPORTED_UPLOAD_FORMATS = {"docx", "pdf"}
SUPPORTED_DOWNLOAD_FORMATS = {"docx", "pdf", "bpmn"}

_ARTIFACT_MIME_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "bpmn": "application/xml",
}


def _safe_filename(value: str, fallback: str) -> str:
    name = Path(value or "").name.strip()
    if not name:
        name = fallback

    name = re.sub(r"[^A-Za-z0-9._()\-\u4e00-\u9fff ]+", "_", name)
    name = name.strip(" .")
    return name[:180] or fallback


def _ensure_storage_directories() -> None:
    UPLOAD_STORAGE_DIR.mkdir(parents=True, exist_ok=True)
    ARTIFACT_STORAGE_DIR.mkdir(parents=True, exist_ok=True)


def _validate_uploaded_file_content(
    *,
    path: Path,
    file_type: str,
) -> None:
    """
    Validate the real file structure instead of trusting the browser MIME type.

    Windows and different browsers may report DOCX files as application/zip,
    application/x-zip-compressed, application/octet-stream, or an empty MIME
    type. The extension plus the actual file signature is more reliable.
    """

    if file_type == "pdf":
        with path.open("rb") as file:
            signature = file.read(5)

        if signature != b"%PDF-":
            raise ValueError(
                "The selected file is not a valid PDF document."
            )

        return

    if file_type == "docx":
        try:
            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
        except zipfile.BadZipFile as exc:
            raise ValueError(
                "The selected file is not a valid DOCX document."
            ) from exc

        required_entries = {
            "[Content_Types].xml",
            "word/document.xml",
        }

        if not required_entries.issubset(names):
            raise ValueError(
                "The selected file is not a valid DOCX document."
            )

        return

    raise ValueError(
        f"Unsupported upload format: {file_type}"
    )


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX support is unavailable. Install python-docx."
        ) from exc

    document = Document(path)
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        table_lines: list[str] = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                table_lines.append(" | ".join(values))
        if table_lines:
            blocks.append(
                f"Table {table_index}:\n" + "\n".join(table_lines)
            )

    return "\n\n".join(blocks).strip()


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "PDF support is unavailable. Install pypdf."
        ) from exc

    reader = PdfReader(str(path))
    pages: list[str] = []

    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"Page {index}:\n{text}")

    return "\n\n".join(pages).strip()


def _extract_text(path: Path, file_type: str) -> str:
    if file_type == "docx":
        return _extract_docx_text(path)
    if file_type == "pdf":
        return _extract_pdf_text(path)
    raise ValueError(f"Unsupported upload format: {file_type}")


def load_workflow_io_configuration(
    *, workflow_id: int, user_id: int
) -> tuple[dict[str, Any], dict[str, Any]]:
    with get_connection() as conn:
        row = conn.execute(
            """
            SELECT
                input_configuration_json,
                output_configuration_json
            FROM workflows
            WHERE id = ?
              AND user_id = ?
            """,
            (workflow_id, user_id),
        ).fetchone()

    if row is None:
        raise HTTPException(status_code=404, detail="Workflow not found.")

    import json

    try:
        input_config = json.loads(row[0] or "{}")
    except json.JSONDecodeError:
        input_config = {}

    try:
        output_config = json.loads(row[1] or "{}")
    except json.JSONDecodeError:
        output_config = {}

    if not isinstance(input_config, dict):
        input_config = {}
    if not isinstance(output_config, dict):
        output_config = {}

    return input_config, output_config


async def store_workflow_upload(
    *, workflow_id: int, user_id: int, uploaded_file: UploadFile
) -> dict[str, Any]:
    input_config, _ = load_workflow_io_configuration(
        workflow_id=workflow_id,
        user_id=user_id,
    )

    upload_config = input_config.get("file_upload")
    if not isinstance(upload_config, dict) or not upload_config.get("enabled"):
        raise HTTPException(
            status_code=400,
            detail="This workflow does not accept file uploads.",
        )

    accepted_formats = {
        str(item).lower()
        for item in upload_config.get("accepted_formats", [])
        if str(item).lower() in SUPPORTED_UPLOAD_FORMATS
    }

    original_filename = _safe_filename(
        uploaded_file.filename or "uploaded-file",
        "uploaded-file",
    )
    extension = Path(original_filename).suffix.lower().lstrip(".")

    if extension not in accepted_formats:
        allowed = ", ".join(sorted(accepted_formats)) or "none"
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed formats: {allowed}.",
        )

    # Do not reject a file only because of the browser-provided MIME type.
    # Edge/Chrome/Windows can report the same DOCX with several MIME values.
    content_type = str(
        uploaded_file.content_type
        or "application/octet-stream"
    )

    content = await uploaded_file.read(MAX_WORKFLOW_UPLOAD_BYTES + 1)
    if not content:
        raise HTTPException(status_code=400, detail="The uploaded file is empty.")
    if len(content) > MAX_WORKFLOW_UPLOAD_BYTES:
        raise HTTPException(
            status_code=413,
            detail=(
                "The uploaded file is too large. "
                f"Maximum size is {MAX_WORKFLOW_UPLOAD_BYTES // (1024 * 1024)} MB."
            ),
        )

    _ensure_storage_directories()
    stored_filename = f"{uuid.uuid4().hex}.{extension}"
    storage_path = UPLOAD_STORAGE_DIR / stored_filename
    storage_path.write_bytes(content)

    try:
        _validate_uploaded_file_content(
            path=storage_path,
            file_type=extension,
        )
        extracted_text = _extract_text(storage_path, extension)
    except Exception as exc:
        storage_path.unlink(missing_ok=True)
        raise HTTPException(
            status_code=400,
            detail=f"Unable to read the uploaded {extension.upper()} file: {exc}",
        ) from exc

    if not extracted_text:
        storage_path.unlink(missing_ok=True)
        raise HTTPException(
            status_code=400,
            detail=(
                "No readable text was found in the uploaded file. "
                "Scanned PDFs are not supported in this version."
            ),
        )

    with get_connection() as conn:
        cursor = conn.execute(
            """
            INSERT INTO workflow_uploaded_files (
                workflow_id,
                user_id,
                original_filename,
                stored_filename,
                file_type,
                mime_type,
                size_bytes,
                file_path,
                extracted_text
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                workflow_id,
                user_id,
                original_filename,
                stored_filename,
                extension,
                content_type,
                len(content),
                str(storage_path),
                extracted_text,
            ),
        )
        conn.commit()

        if cursor.lastrowid is None:
            storage_path.unlink(missing_ok=True)
            raise HTTPException(
                status_code=500,
                detail="Unable to save the uploaded file.",
            )

        file_id = int(cursor.lastrowid)

    return {
        "id": file_id,
        "workflow_id": workflow_id,
        "filename": original_filename,
        "file_type": extension,
        "mime_type": content_type,
        "size_bytes": len(content),
        "text_preview": extracted_text[:500],
    }


def delete_workflow_upload(
    *, workflow_id: int, user_id: int, file_id: int
) -> None:
    with get_connection() as conn:
        row = conn.execute(
            """
            SELECT file_path
            FROM workflow_uploaded_files
            WHERE id = ?
              AND workflow_id = ?
              AND user_id = ?
            """,
            (file_id, workflow_id, user_id),
        ).fetchone()

        if row is None:
            raise HTTPException(status_code=404, detail="Uploaded file not found.")

        conn.execute(
            """
            DELETE FROM workflow_uploaded_files
            WHERE id = ?
              AND workflow_id = ?
              AND user_id = ?
            """,
            (file_id, workflow_id, user_id),
        )
        conn.commit()

    Path(str(row[0])).unlink(missing_ok=True)


def build_workflow_input_from_files(
    *,
    workflow_id: int,
    user_id: int,
    message: str,
    file_ids: list[int],
) -> tuple[str, str]:
    normalized_ids = list(dict.fromkeys(int(item) for item in file_ids))

    if not normalized_ids:
        clean_message = message.strip()
        return clean_message, clean_message

    input_config, _ = load_workflow_io_configuration(
        workflow_id=workflow_id,
        user_id=user_id,
    )
    upload_config = input_config.get("file_upload")

    if not isinstance(upload_config, dict) or not upload_config.get("enabled"):
        raise HTTPException(
            status_code=400,
            detail="This workflow does not accept file uploads.",
        )

    max_files = int(upload_config.get("max_files") or 1)
    if len(normalized_ids) > max_files:
        raise HTTPException(
            status_code=400,
            detail=f"This workflow accepts at most {max_files} uploaded file(s).",
        )

    placeholders = ",".join("?" for _ in normalized_ids)

    with get_connection() as conn:
        rows = conn.execute(
            f"""
            SELECT
                id,
                original_filename,
                file_type,
                extracted_text
            FROM workflow_uploaded_files
            WHERE workflow_id = ?
              AND user_id = ?
              AND id IN ({placeholders})
            ORDER BY id ASC
            """,
            (workflow_id, user_id, *normalized_ids),
        ).fetchall()

    if len(rows) != len(normalized_ids):
        raise HTTPException(
            status_code=404,
            detail="One or more uploaded files could not be found.",
        )

    sections: list[str] = []
    display_names: list[str] = []

    for row in rows:
        filename = str(row[1])
        file_type = str(row[2]).upper()
        extracted_text = str(row[3] or "").strip()
        display_names.append(filename)
        sections.append(
            f"Uploaded file: {filename}\n"
            f"Format: {file_type}\n\n"
            f"Extracted content:\n{extracted_text}"
        )

    clean_message = message.strip()
    instruction = clean_message or (
        "Analyze the uploaded document and complete this workflow's assigned task."
    )

    runtime_input = (
        f"User request:\n{instruction}\n\n"
        "Uploaded documents:\n\n"
        + "\n\n---\n\n".join(sections)
    )

    display_input = instruction + "\n\nFiles: " + ", ".join(display_names)
    return runtime_input, display_input


def _markdown_to_plain_lines(text: str) -> list[tuple[str, str]]:
    """
    Convert stored plain text into presentation roles.

    The stored content remains Markdown-free. Lines such as "Summary:"
    and "Action Items:" are treated as visual headings only when exporting.
    """

    result: list[tuple[str, str]] = []

    for raw_line in text.splitlines():
        line = raw_line.strip()

        if not line:
            result.append(("blank", ""))
        elif line.startswith("### "):
            result.append(("h3", line[4:].strip()))
        elif line.startswith("## "):
            result.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            result.append(("h1", line[2:].strip()))
        elif (
            line.endswith((":",
                           "："))
            and len(line) <= 100
            and not re.match(r"^\d+[.)]\s+", line)
        ):
            result.append(("h2", line))
        elif re.match(r"^\s*[•●▪*-]\s+", line):
            result.append(
                (
                    "bullet",
                    re.sub(
                        r"^\s*[•●▪*-]\s+",
                        "",
                        line,
                    ),
                )
            )
        elif re.match(r"^\s*\d+[.)]\s+", line):
            result.append(
                (
                    "number",
                    re.sub(
                        r"^\s*\d+[.)]\s+",
                        "",
                        line,
                    ),
                )
            )
        else:
            cleaned = re.sub(
                r"\*\*(.*?)\*\*",
                r"\1",
                line,
            )
            cleaned = re.sub(
                r"`([^`]*)`",
                r"\1",
                cleaned,
            )
            result.append(
                ("paragraph", cleaned)
            )

    return result


def _generate_docx(path: Path, title: str, final_output: str) -> None:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Install python-docx to generate DOCX files.") from exc

    document = Document()
    document.add_heading(title, level=0)

    for kind, value in _markdown_to_plain_lines(final_output):
        if kind == "blank":
            document.add_paragraph("")
        elif kind == "h1":
            document.add_heading(value, level=1)
        elif kind == "h2":
            document.add_heading(value, level=2)
        elif kind == "h3":
            document.add_heading(value, level=3)
        elif kind == "bullet":
            document.add_paragraph(value, style="List Bullet")
        elif kind == "number":
            document.add_paragraph(value, style="List Number")
        else:
            document.add_paragraph(value)

    document.save(path)


def _find_unicode_font() -> str | None:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/msyh.ttf"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("/System/Library/Fonts/PingFang.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return None


def _generate_pdf(path: Path, title: str, final_output: str) -> None:
    try:
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("Install reportlab to generate PDF files.") from exc

    font_name = "Helvetica"
    font_path = _find_unicode_font()
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("WorkflowUnicode", font_path))
            font_name = "WorkflowUnicode"
        except Exception:
            font_name = "Helvetica"

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "WorkflowTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=20,
        leading=26,
        spaceAfter=16,
    )
    body_style = ParagraphStyle(
        "WorkflowBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=16,
        alignment=TA_LEFT,
        spaceAfter=7,
    )
    heading_styles = {
        "h1": ParagraphStyle("H1", parent=body_style, fontSize=16, leading=21, spaceBefore=12, spaceAfter=8),
        "h2": ParagraphStyle("H2", parent=body_style, fontSize=14, leading=19, spaceBefore=10, spaceAfter=7),
        "h3": ParagraphStyle("H3", parent=body_style, fontSize=12, leading=17, spaceBefore=8, spaceAfter=6),
    }

    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        title=title,
    )

    story: list[Any] = [Paragraph(html.escape(title), title_style)]
    number_index = 0

    for kind, value in _markdown_to_plain_lines(final_output):
        if kind == "blank":
            story.append(Spacer(1, 4))
            continue
        if kind in heading_styles:
            story.append(Paragraph(html.escape(value), heading_styles[kind]))
            continue
        if kind == "bullet":
            value = f"• {value}"
        elif kind == "number":
            number_index += 1
            value = f"{number_index}. {value}"
        else:
            number_index = 0

        safe_value = html.escape(value).replace("\n", "<br/>")
        try:
            story.append(Paragraph(safe_value, body_style))
        except UnicodeEncodeError:
            fallback = value.encode("latin-1", errors="replace").decode("latin-1")
            story.append(Paragraph(html.escape(fallback), body_style))

    document.build(story)


def _generate_bpmn(path: Path, workflow_name: str, agents: list[dict[str, Any]]) -> None:
    definitions = ET.Element(
        "bpmn:definitions",
        {
            "xmlns:bpmn": "http://www.omg.org/spec/BPMN/20100524/MODEL",
            "xmlns:bpmndi": "http://www.omg.org/spec/BPMN/20100524/DI",
            "xmlns:dc": "http://www.omg.org/spec/DD/20100524/DC",
            "xmlns:di": "http://www.omg.org/spec/DD/20100524/DI",
            "id": "Definitions_1",
            "targetNamespace": "http://agentdemo.local/bpmn",
        },
    )

    process_id = f"Process_{uuid.uuid4().hex[:8]}"
    process = ET.SubElement(
        definitions,
        "bpmn:process",
        {"id": process_id, "name": workflow_name, "isExecutable": "false"},
    )

    node_ids = ["StartEvent_1"]
    ET.SubElement(process, "bpmn:startEvent", {"id": "StartEvent_1", "name": "Start"})

    for index, agent in enumerate(agents, start=1):
        node_id = f"Task_{index}"
        node_ids.append(node_id)
        ET.SubElement(
            process,
            "bpmn:task",
            {"id": node_id, "name": str(agent.get("name") or f"Agent {index}")},
        )

    node_ids.append("EndEvent_1")
    ET.SubElement(process, "bpmn:endEvent", {"id": "EndEvent_1", "name": "End"})

    flow_ids: list[str] = []
    for index in range(len(node_ids) - 1):
        flow_id = f"Flow_{index + 1}"
        flow_ids.append(flow_id)
        ET.SubElement(
            process,
            "bpmn:sequenceFlow",
            {
                "id": flow_id,
                "sourceRef": node_ids[index],
                "targetRef": node_ids[index + 1],
            },
        )

    diagram = ET.SubElement(definitions, "bpmndi:BPMNDiagram", {"id": "BPMNDiagram_1"})
    plane = ET.SubElement(
        diagram,
        "bpmndi:BPMNPlane",
        {"id": "BPMNPlane_1", "bpmnElement": process_id},
    )

    positions: dict[str, tuple[int, int, int, int]] = {}
    x = 100
    for node_id in node_ids:
        if node_id.startswith("Start") or node_id.startswith("End"):
            width, height = 36, 36
            y = 122
        else:
            width, height = 130, 80
            y = 100
        positions[node_id] = (x, y, width, height)
        shape = ET.SubElement(
            plane,
            "bpmndi:BPMNShape",
            {"id": f"{node_id}_di", "bpmnElement": node_id},
        )
        ET.SubElement(
            shape,
            "dc:Bounds",
            {"x": str(x), "y": str(y), "width": str(width), "height": str(height)},
        )
        x += width + 90

    for index, flow_id in enumerate(flow_ids):
        source_id = node_ids[index]
        target_id = node_ids[index + 1]
        sx, sy, sw, sh = positions[source_id]
        tx, ty, tw, th = positions[target_id]
        edge = ET.SubElement(
            plane,
            "bpmndi:BPMNEdge",
            {"id": f"{flow_id}_di", "bpmnElement": flow_id},
        )
        ET.SubElement(
            edge,
            "di:waypoint",
            {"x": str(sx + sw), "y": str(sy + sh // 2)},
        )
        ET.SubElement(
            edge,
            "di:waypoint",
            {"x": str(tx), "y": str(ty + th // 2)},
        )

    tree = ET.ElementTree(definitions)
    ET.indent(tree, space="  ")
    tree.write(path, encoding="utf-8", xml_declaration=True)


def artifact_row_to_dict(row: Any) -> dict[str, Any]:
    return {
        "id": int(row[0]),
        "run_id": int(row[1]),
        "artifact_type": str(row[2]),
        "filename": str(row[3]),
        "mime_type": str(row[4]),
        "download_url": f"/api/workflow-artifacts/{int(row[0])}/download",
        "created_at": str(row[5]),
    }


def load_run_artifacts(
    *, run_id: int, workflow_id: int, user_id: int
) -> list[dict[str, Any]]:
    with get_connection() as conn:
        rows = conn.execute(
            """
            SELECT
                id,
                run_id,
                artifact_type,
                filename,
                mime_type,
                created_at
            FROM workflow_run_artifacts
            WHERE run_id = ?
              AND workflow_id = ?
              AND user_id = ?
            ORDER BY id ASC
            """,
            (run_id, workflow_id, user_id),
        ).fetchall()

    return [artifact_row_to_dict(row) for row in rows]


def generate_run_artifacts(
    *,
    run_id: int,
    workflow_id: int,
    user_id: int,
    workflow_name: str,
    agents: list[dict[str, Any]],
    final_output: str,
    formats: list[str],
) -> list[dict[str, Any]]:
    normalized_formats = [
        item
        for item in dict.fromkeys(str(value).lower() for value in formats)
        if item in SUPPORTED_DOWNLOAD_FORMATS
    ]

    if not normalized_formats:
        return []

    _ensure_storage_directories()
    created_ids: list[int] = []

    for artifact_type in normalized_formats:
        base_name = _safe_filename(workflow_name, "workflow-result")
        filename = f"{base_name}-run-{run_id}.{artifact_type}"
        stored_name = f"{uuid.uuid4().hex}.{artifact_type}"
        path = ARTIFACT_STORAGE_DIR / stored_name

        if artifact_type == "docx":
            _generate_docx(path, workflow_name, final_output)
        elif artifact_type == "pdf":
            _generate_pdf(path, workflow_name, final_output)
        elif artifact_type == "bpmn":
            _generate_bpmn(path, workflow_name, agents)

        with get_connection() as conn:
            cursor = conn.execute(
                """
                INSERT INTO workflow_run_artifacts (
                    run_id,
                    workflow_id,
                    user_id,
                    artifact_type,
                    filename,
                    stored_filename,
                    file_path,
                    mime_type
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    run_id,
                    workflow_id,
                    user_id,
                    artifact_type,
                    filename,
                    stored_name,
                    str(path),
                    _ARTIFACT_MIME_TYPES[artifact_type],
                ),
            )
            conn.commit()
            if cursor.lastrowid is not None:
                created_ids.append(int(cursor.lastrowid))

    return load_run_artifacts(
        run_id=run_id,
        workflow_id=workflow_id,
        user_id=user_id,
    )


def get_artifact_download(
    *, artifact_id: int, user_id: int
) -> tuple[Path, str, str]:
    with get_connection() as conn:
        row = conn.execute(
            """
            SELECT file_path, filename, mime_type
            FROM workflow_run_artifacts
            WHERE id = ?
              AND user_id = ?
            """,
            (artifact_id, user_id),
        ).fetchone()

    if row is None:
        raise HTTPException(status_code=404, detail="Generated file not found.")

    path = Path(str(row[0]))
    if not path.exists() or not path.is_file():
        raise HTTPException(
            status_code=404,
            detail="The generated file is no longer available.",
        )

    return path, str(row[1]), str(row[2])
